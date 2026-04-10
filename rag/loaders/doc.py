import importlib
from tqdm import tqdm
from typing import Literal
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from unstructured.partition.text import partition_text
from langchain_community.document_loaders import UnstructuredFileLoader

from rag.loaders.img import IMAGE_LOADER_METHODS


class DOCLoader(UnstructuredFileLoader):
    def __init__(self, file_path, method: Literal["rapidocr", "llmcaption", "blipcaption"]="rapidocr", **kwargs):
        super().__init__(file_path, **kwargs)
        module = importlib.import_module("rag.loaders.img")
        ImageLoader = getattr(module, IMAGE_LOADER_METHODS[method])
        self.image_loader = ImageLoader()
    def para2text(self, para: Paragraph):
        para_text = para.text + "\n"
        image_ids = para._element.xpath(".//a:blip/@r:embed")
        with tqdm(total=len(image_ids), position=0, leave=False,
                desc="DOC Loading With Image Processing") as pbar2:
            for image_id in image_ids:
                image = para.part.related_parts[image_id].image
                if image.px_width > 10 and image.px_height > 10:
                    if image_content := self.image_loader.img2text(image.blob):
                        image_content = image_content.replace("\n", " ")
                        para_text += f"这里有一张图片内容为：“{image_content}”。\n"
                pbar2.update()
        return para_text
    def doc2text(self, file_path):
        content = ""
        doc = Document(file_path)
        total = len(doc.paragraphs) + len(doc.tables)
        with tqdm(total=total, position=0, leave=True, desc="DOC Loading") as pbar:
            for para_or_table in doc.iter_inner_content():
                if isinstance(para_or_table, Paragraph):
                    content += self.para2text(para_or_table)
                elif isinstance(para_or_table, Table):
                    table_content = ""
                    for row in para_or_table.rows:
                        table_content += "| "
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                para_text = self.para2text(para)
                                table_content += para_text.replace("\n", " ")
                            table_content += "| "
                        table_content += "\n"
                    content += f"这里有一个表格内容为：\n{table_content}表格至此结束。\n"
                pbar.update()
        return content
    def _get_elements(self):
        text = self.doc2text(self.file_path)
        return partition_text(text=text, **self.unstructured_kwargs)
